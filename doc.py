from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# Créer le document
doc = Document()

# Configuration des marges
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Fonction pour ajouter une bordure à un paragraphe
def add_border_to_paragraph(paragraph, border_color="FF6600"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:color'), border_color)
        pBdr.append(border)
    pPr.append(pBdr)

# Fonction pour créer un tableau stylisé
def set_table_style(table, header_color="FF6600"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
            if i == 0:  # Header row
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), header_color)
                cell._tc.get_or_add_tcPr().append(shading)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True

# ==================== PAGE DE GARDE ====================

# Logo/Titre principal
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("COMPAGNIE IVOIRIENNE D'ÉLECTRICITÉ")
title_run.bold = True
title_run.font.size = Pt(28)
title_run.font.color.rgb = RGBColor(255, 102, 0)

# Slogan
slogan = doc.add_paragraph()
slogan.alignment = WD_ALIGN_PARAGRAPH.CENTER
slogan_run = slogan.add_run("L'énergie au service du développement")
slogan_run.italic = True
slogan_run.font.size = Pt(14)
slogan_run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
doc.add_paragraph()

# Titre du document
main_title = doc.add_paragraph()
main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
main_title_run = main_title.add_run("CAHIER DES CHARGES")
main_title_run.bold = True
main_title_run.font.size = Pt(36)
main_title_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

# Sous-titre
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("Plateforme Digitale de Services & Souscription")
subtitle_run.font.size = Pt(18)
subtitle_run.font.color.rgb = RGBColor(0, 102, 153)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Cadre projet
project_box = doc.add_paragraph()
project_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_border_to_paragraph(project_box, "003366")
project_box_run = project_box.add_run("\nPROJET DE DÉVELOPPEMENT WEB\n\n")
project_box_run.bold = True
project_box_run.font.size = Pt(16)
project_box_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
doc.add_paragraph()

# Informations du document
info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_para.add_run("Version : 1.0\n").font.size = Pt(12)
info_para.add_run(f"Date : {datetime.datetime.now().strftime('%d/%m/%Y')}\n").font.size = Pt(12)
info_para.add_run("Statut : En attente de validation\n").font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

# Bas de page de garde
footer_info = doc.add_paragraph()
footer_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_info.add_run("Document confidentiel - Tous droits réservés")
footer_run.italic = True
footer_run.font.size = Pt(10)
footer_run.font.color.rgb = RGBColor(128, 128, 128)

# Saut de page
doc.add_page_break()

# ==================== SOMMAIRE ====================

toc_title = doc.add_paragraph()
toc_title_run = toc_title.add_run("SOMMAIRE")
toc_title_run.bold = True
toc_title_run.font.size = Pt(20)
toc_title_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

toc_items = [
    ("1.", "Lettre de présentation", "3"),
    ("2.", "Contexte et enjeux", "4"),
    ("3.", "Objectifs du projet", "6"),
    ("4.", "Périmètre fonctionnel", "7"),
    ("5.", "Architecture technique", "15"),
    ("6.", "Design et expérience utilisateur", "17"),
    ("7.", "Planification", "18"),
    ("8.", "Budget prévisionnel", "19"),
    ("9.", "Équipe projet", "20"),
    ("10.", "Risques et mitigation", "21"),
    ("11.", "Annexes", "22"),
]

for num, title_text, page in toc_items:
    toc_line = doc.add_paragraph()
    toc_line.add_run(f"{num} {title_text}").font.size = Pt(12)
    toc_line.add_run(f"{'.' * (60 - len(title_text))}{page}").font.size = Pt(12)

doc.add_page_break()

# ==================== 1. LETTRE DE PRÉSENTATION ====================

h1 = doc.add_paragraph()
h1_run = h1.add_run("1. LETTRE DE PRÉSENTATION")
h1_run.bold = True
h1_run.font.size = Pt(18)
h1_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

# Cadre destinataire
dest = doc.add_paragraph()
dest.add_run("À l'attention de la Direction Générale\n").bold = True
dest.add_run("Compagnie Ivoirienne d'Electricité\n")
dest.add_run("Abidjan, Côte d'Ivoire")

doc.add_paragraph()

# Objet
obj = doc.add_paragraph()
obj.add_run("Objet : ").bold = True
obj.add_run("Proposition de développement d'une plateforme digitale de services aux usagers")

doc.add_paragraph()

# Corps de la lettre
letter_paragraphs = [
    "Madame, Monsieur le Directeur Général,",
    "C'est avec un grand honneur que nous vous soumettons le présent cahier de charges pour le développement d'une plateforme digitale innovante au service de votre stratégie de transformation numérique.",
    "La CIE, acteur historique et incontournable du secteur énergétique ivoirien, a démontré son engagement envers la population à travers des programmes ambitieux tels que le Programme Électricité Pour Tous (PEPT). Notre proposition vise à amplifier cet impact en digitalisant le parcours client, de la découverte des offres jusqu'au suivi des demandes de raccordement.",
    "Cette solution répond à une double ambition :\n• Rapprocher la CIE de ses usagers grâce à une présence digitale moderne et accessible\n• Optimiser les processus internes par la digitalisation du traitement des demandes",
    "Nous restons à votre entière disposition pour présenter ce projet en détail et adapter notre proposition à vos spécificités organisationnelles.",
    "Nous vous prions d'agréer, Madame, Monsieur le Directeur Général, l'expression de notre haute considération."
]

for para_text in letter_paragraphs:
    p = doc.add_paragraph()
    p.add_run(para_text).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(12)

doc.add_paragraph()

# Signature
sig_table = doc.add_table(rows=2, cols=2)
sig_table.autosize = True
sig_table.rows[0].cells[0].paragraphs[0].add_run("Le Porteur du Projet\n\n_____________________").font.size = Pt(11)
sig_table.rows[0].cells[1].paragraphs[0].add_run(f"Date : {datetime.datetime.now().strftime('%d/%m/%Y')}\nLieu : Abidjan").font.size = Pt(11)

doc.add_page_break()

# ==================== 2. CONTEXTE ET ENJEUX ====================

h2 = doc.add_paragraph()
h2_run = h2.add_run("2. CONTEXTE ET ENJEUX")
h2_run.bold = True
h2_run.font.size = Pt(18)
h2_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

# 2.1 Présentation de la CIE
h2_1 = doc.add_paragraph()
h2_1_run = h2_1.add_run("2.1 Présentation de la CIE")
h2_1_run.bold = True
h2_1_run.font.size = Pt(14)
h2_1_run.font.color.rgb = RGBColor(255, 102, 0)

p = doc.add_paragraph()
p.add_run("""La Compagnie Ivoirienne d'Electricité (CIE) est le pilier du secteur énergétique en Côte d'Ivoire. Elle assure la production, le transport et la distribution d'électricité sur l'ensemble du territoire national, avec pour mission principale l'électrification du pays et l'accompagnement des populations vers l'accès à l'énergie.""").font.size = Pt(11)

doc.add_paragraph()

# 2.2 Constat actuel
h2_2 = doc.add_paragraph()
h2_2_run = h2_2.add_run("2.2 Constat actuel")
h2_2_run.bold = True
h2_2_run.font.size = Pt(14)
h2_2_run.font.color.rgb = RGBColor(255, 102, 0)

table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
headers = ["Aspect", "Situation actuelle", "Impact"]
for i, header in enumerate(headers):
    table.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

data = [
    ("Souscription", "Processus papier, déplacements physiques", "Perte de temps, coûts logistiques"),
    ("Information", "Canaux dispersés, peu accessibles", "Confusion, manque de visibilité"),
    ("Suivi", "Aucun outil de tracking", "Appels récurrents, insatisfaction"),
    ("Traitement", "Processus manuels", "Délais importants, erreurs possibles"),
]

for i, (aspect, situation, impact) in enumerate(data, 1):
    table.rows[i].cells[0].paragraphs[0].add_run(aspect).font.size = Pt(10)
    table.rows[i].cells[1].paragraphs[0].add_run(situation).font.size = Pt(10)
    table.rows[i].cells[2].paragraphs[0].add_run(impact).font.size = Pt(10)

set_table_style(table)

doc.add_paragraph()

# 2.3 Programme PEPT
h2_3 = doc.add_paragraph()
h2_3_run = h2_3.add_run("2.3 Programme Électricité Pour Tous (PEPT)")
h2_3_run.bold = True
h2_3_run.font.size = Pt(14)
h2_3_run.font.color.rgb = RGBColor(255, 102, 0)

p = doc.add_paragraph()
p.add_run("Le Programme Électricité Pour Tous est une initiative majeure de la CIE pour démocratiser l'accès à l'électricité :").font.size = Pt(11)

pept_items = [
    "Souscription GRATUITE",
    "Installation kit appartement (1-3 pièces)",
    "Raccordement + Abonnement inclus",
    "2 kWh d'énergie OFFERTS à l'activation",
    "Sans caution ni dépôt de garantie",
    "Crédit de raccordement remboursable sur 5 ans"
]

for item in pept_items:
    p = doc.add_paragraph()
    p.add_run(f"✓ {item}").font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(1)

doc.add_paragraph()

# Tableau des offres PEPT
p = doc.add_paragraph()
p.add_run("Offres tarifaires PEPT :").bold = True

table2 = doc.add_table(rows=3, cols=4)
table2.style = 'Table Grid'
headers2 = ["Type d'offre", "Zone", "Prix (FCFA)", "Caractéristiques"]
for i, header in enumerate(headers2):
    table2.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

data2 = [
    ("Offre Standard", "Centres ruraux", "85 000", "Compteur monophasé 10(40)A"),
    ("Offre Standard", "Grandes villes", "150 000", "Compteur monophasé 10(40)A"),
]

for i, row_data in enumerate(data2, 1):
    for j, cell_data in enumerate(row_data):
        table2.rows[i].cells[j].paragraphs[0].add_run(cell_data).font.size = Pt(10)

set_table_style(table2)

doc.add_page_break()

# ==================== 3. OBJECTIFS DU PROJET ====================

h3 = doc.add_paragraph()
h3_run = h3.add_run("3. OBJECTIFS DU PROJET")
h3_run.bold = True
h3_run.font.size = Pt(18)
h3_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

h3_1 = doc.add_paragraph()
h3_1_run = h3_1.add_run("3.1 Objectif général")
h3_1_run.bold = True
h3_1_run.font.size = Pt(14)
h3_1_run.font.color.rgb = RGBColor(255, 102, 0)

p = doc.add_paragraph()
p.add_run("Développer une plateforme web intégrée permettant de présenter les services de la CIE et de digitaliser le processus de souscription au Programme Électricité Pour Tous.").font.size = Pt(11)

doc.add_paragraph()

h3_2 = doc.add_paragraph()
h3_2_run = h3_2.add_run("3.2 Objectifs spécifiques")
h3_2_run.bold = True
h3_2_run.font.size = Pt(14)
h3_2_run.font.color.rgb = RGBColor(255, 102, 0)

table3 = doc.add_table(rows=6, cols=4)
table3.style = 'Table Grid'
headers3 = ["N°", "Objectif", "Indicateur de succès", "Cible"]
for i, header in enumerate(headers3):
    table3.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

objectives = [
    ("1", "Digitaliser l'information sur les services", "Taux de visiteurs informés", "80%"),
    ("2", "Automatiser les souscriptions", "Nombre de demandes en ligne", "1000/mois"),
    ("3", "Assurer le suivi des demandes", "Taux de suivi effectué", "90%"),
    ("4", "Optimiser le traitement interne", "Délai moyen de traitement", "-40%"),
    ("5", "Améliorer la satisfaction client", "Score de satisfaction", "4/5"),
]

for i, (num, obj_text, indic, cible) in enumerate(objectives, 1):
    table3.rows[i].cells[0].paragraphs[0].add_run(num).font.size = Pt(10)
    table3.rows[i].cells[1].paragraphs[0].add_run(obj_text).font.size = Pt(10)
    table3.rows[i].cells[2].paragraphs[0].add_run(indic).font.size = Pt(10)
    table3.rows[i].cells[3].paragraphs[0].add_run(cible).font.size = Pt(10)

set_table_style(table3)

doc.add_paragraph()

h3_3 = doc.add_paragraph()
h3_3_run = h3_3.add_run("3.3 Périmètre géographique")
h3_3_run.bold = True
h3_3_run.font.size = Pt(14)
h3_3_run.font.color.rgb = RGBColor(255, 102, 0)

phases = [
    ("PHASE 1", "Abidjan et périphérie"),
    ("PHASE 2", "Grandes villes intérieures (Yamoussoukro, Bouaké, San-Pédro, Daloa...)"),
    ("PHASE 3", "Couverture nationale complète"),
]

for phase, desc in phases:
    p = doc.add_paragraph()
    run1 = p.add_run(f"{phase} : ")
    run1.bold = True
    run1.font.size = Pt(11)
    p.add_run(desc).font.size = Pt(11)

doc.add_page_break()

# ==================== 4. PÉRIMÈTRE FONCTIONNEL ====================

h4 = doc.add_paragraph()
h4_run = h4.add_run("4. PÉRIMÈTRE FONCTIONNEL")
h4_run.bold = True
h4_run.font.size = Pt(18)
h4_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

# 4.1 Architecture globale
h4_1 = doc.add_paragraph()
h4_1_run = h4_1.add_run("4.1 Architecture fonctionnelle globale")
h4_1_run.bold = True
h4_1_run.font.size = Pt(14)
h4_1_run.font.color.rgb = RGBColor(255, 102, 0)

p = doc.add_paragraph()
p.add_run("La plateforme se compose de trois interfaces principales :").font.size = Pt(11)

interfaces = [
    "Interface Usager : Site vitrine, formulaire de souscription, suivi de demande",
    "Backend / API : Gestion des demandes, génération de codes, notifications",
    "Interface Agent CIE : Tableau de bord, traitement des demandes, statistiques"
]

for interface in interfaces:
    p = doc.add_paragraph()
    p.add_run(f"• {interface}").font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(1)

doc.add_paragraph()

# 4.2 Module 1 : Site vitrine
h4_2 = doc.add_paragraph()
h4_2_run = h4_2.add_run("4.2 Module 1 : Site vitrine de présentation des services")
h4_2_run.bold = True
h4_2_run.font.size = Pt(14)
h4_2_run.font.color.rgb = RGBColor(255, 102, 0)

p = doc.add_paragraph()
run_pages = p.add_run("Pages du site :")
run_pages.bold = True
run_pages.font.size = Pt(11)

pages = [
    ("Page d'accueil", "Présentation générale, bannière, CTAs, actualités, témoignages"),
    ("Page Services", "Liste des services CIE : raccordement, facturation, assistance"),
    ("Page PEPT", "Présentation détaillée du programme, offres, avantages, FAQ"),
    ("Page Contact", "Coordonnées, formulaire de contact, localisation agences"),
]

for page, desc in pages:
    p = doc.add_paragraph()
    run_page = p.add_run(f"• {page} : ")
    run_page.bold = True
    run_page.font.size = Pt(11)
    p.add_run(desc).font.size = Pt(11)

doc.add_paragraph()

# 4.3 Module 2 : Formulaire de souscription
h4_3 = doc.add_paragraph()
h4_3_run = h4_3.add_run("4.3 Module 2 : Formulaire de souscription PEPT")
h4_3_run.bold = True
h4_3_run.font.size = Pt(14)
h4_3_run.font.color.rgb = RGBColor(255, 102, 0)

p = doc.add_paragraph()
p.add_run("Parcours en 4 étapes :").font.size = Pt(11)

etapes = [
    ("Étape 1", "Choix de l'offre (Standard/Premium)"),
    ("Étape 2", "Saisie des informations personnelles et adresse"),
    ("Étape 3", "Confirmation et récapitulatif"),
    ("Étape 4", "Génération et envoi du code de suivi"),
]

for etape, desc in etapes:
    p = doc.add_paragraph()
    run_etape = p.add_run(f"{etape} : ")
    run_etape.bold = True
    run_etape.font.size = Pt(11)
    p.add_run(desc).font.size = Pt(11)

doc.add_paragraph()

p = doc.add_paragraph()
run_champs = p.add_run("Champs du formulaire :")
run_champs.bold = True
run_champs.font.size = Pt(11)

# Tableau des champs
table4 = doc.add_table(rows=6, cols=3)
table4.style = 'Table Grid'
headers4 = ["Section", "Champs", "Obligatoire"]
for i, header in enumerate(headers4):
    table4.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

champs = [
    ("Informations personnelles", "Civilité, Nom, Prénoms, Date naissance, N° CNI, Téléphone, Email", "Oui"),
    ("Adresse raccordement", "Région, Département, Commune, Quartier, Rue, Point de repère", "Oui"),
    ("Type installation", "Type logement, Nombre pièces, Usage (domestique/professionnel)", "Oui"),
    ("Choix offre", "Zone (rurale/urbaine), Type compteur (mono/triphasé)", "Oui"),
    ("Documents", "Pièce d'identité (recto/verso), Plan de situation", "Recommandé"),
]

for i, (section, champs_list, oblig) in enumerate(champs, 1):
    table4.rows[i].cells[0].paragraphs[0].add_run(section).font.size = Pt(10)
    table4.rows[i].cells[1].paragraphs[0].add_run(champs_list).font.size = Pt(10)
    table4.rows[i].cells[2].paragraphs[0].add_run(oblig).font.size = Pt(10)

set_table_style(table4)

doc.add_paragraph()

# 4.4 Code de suivi
h4_4 = doc.add_paragraph()
h4_4_run = h4_4.add_run("4.4 Génération du code de suivi")
h4_4_run.bold = True
h4_4_run.font.size = Pt(14)
h4_4_run.font.color.rgb = RGBColor(255, 102, 0)

p = doc.add_paragraph()
run_code = p.add_run("Format du code : ")
run_code.bold = True
run_code.font.size = Pt(11)
p.add_run("CIE-YYYY-XXXXXXXX (ex: CIE-2025-A7F3B9K2)").font.size = Pt(11)

p = doc.add_paragraph()
p.add_run("Canaux d'envoi : SMS, Email").font.size = Pt(11)

doc.add_paragraph()

# 4.5 Module 3 : Page de suivi
h4_5 = doc.add_paragraph()
h4_5_run = h4_5.add_run("4.5 Module 3 : Page de suivi de demande")
h4_5_run.bold = True
h4_5_run.font.size = Pt(14)
h4_5_run.font.color.rgb = RGBColor(255, 102, 0)

p = doc.add_paragraph()
p.add_run("Fonctionnalités :").font.size = Pt(11)

suivi_features = [
    "Saisie du code de suivi",
    "Affichage du statut en temps réel",
    "Historique des actions",
    "Informations de contact en cas de besoin"
]

for feature in suivi_features:
    p = doc.add_paragraph()
    p.add_run(f"• {feature}").font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(1)

doc.add_paragraph()

p = doc.add_paragraph()
run_statuts = p.add_run("Statuts possibles :")
run_statuts.bold = True
run_statuts.font.size = Pt(11)

table5 = doc.add_table(rows=7, cols=3)
table5.style = 'Table Grid'
headers5 = ["Statut", "Description", "Action requise"]
for i, header in enumerate(headers5):
    table5.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

statuts = [
    ("EN_ATTENTE", "Demande reçue, en attente de traitement", "Aucune"),
    ("EN_COURS", "Demande en cours d'analyse", "Aucune"),
    ("DOCUMENTS_REQUIS", "Documents complémentaires nécessaires", "Fournir documents"),
    ("PLANIFIEE", "Intervention technique planifiée", "Confirmer date"),
    ("REALISEE", "Raccordement effectué", "Activer compteur"),
    ("REJETEE", "Demande non conforme", "Corriger et resoumettre"),
]

for i, (statut, desc, action) in enumerate(statuts, 1):
    table5.rows[i].cells[0].paragraphs[0].add_run(statut).font.size = Pt(10)
    table5.rows[i].cells[1].paragraphs[0].add_run(desc).font.size = Pt(10)
    table5.rows[i].cells[2].paragraphs[0].add_run(action).font.size = Pt(10)

set_table_style(table5)

doc.add_page_break()

# 4.6 Module 4 : Interface Agent
h4_6 = doc.add_paragraph()
h4_6_run = h4_6.add_run("4.6 Module 4 : Interface Agent CIE")
h4_6_run.bold = True
h4_6_run.font.size = Pt(14)
h4_6_run.font.color.rgb = RGBColor(255, 102, 0)

p = doc.add_paragraph()
p.add_run("Cette interface sécurisée est réservée aux agents de la CIE pour le traitement des demandes.").font.size = Pt(11)

doc.add_paragraph()

p = doc.add_paragraph()
run_agent = p.add_run("Fonctionnalités principales :")
run_agent.bold = True
run_agent.font.size = Pt(11)

agent_features = [
    "Authentification sécurisée (login/mot de passe)",
    "Tableau de bord avec indicateurs clés",
    "Liste des demandes avec filtres (statut, date, région)",
    "Vue détaillée de chaque demande",
    "Actions de traitement (valider, rejeter, demander complément)",
    "Assignation des demandes aux techniciens",
    "Statistiques et rapports exportables"
]

for feature in agent_features:
    p = doc.add_paragraph()
    p.add_run(f"• {feature}").font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(1)

doc.add_paragraph()

p = doc.add_paragraph()
run_kpi = p.add_run("Tableau de bord - Indicateurs :")
run_kpi.bold = True
run_kpi.font.size = Pt(11)

table6 = doc.add_table(rows=5, cols=2)
table6.style = 'Table Grid'
headers6 = ["Indicateur", "Description"]
for i, header in enumerate(headers6):
    table6.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

kpis = [
    ("Demandes totales", "Nombre cumulé de demandes"),
    ("Demandes en attente", "Demandes non traitées"),
    ("Délai moyen de traitement", "Temps entre soumission et résolution"),
    ("Taux de satisfaction", "Basé sur les retours clients"),
]

for i, (ind, desc) in enumerate(kpis, 1):
    table6.rows[i].cells[0].paragraphs[0].add_run(ind).font.size = Pt(10)
    table6.rows[i].cells[1].paragraphs[0].add_run(desc).font.size = Pt(10)

set_table_style(table6)

doc.add_page_break()

# ==================== 5. ARCHITECTURE TECHNIQUE ====================

h5 = doc.add_paragraph()
h5_run = h5.add_run("5. ARCHITECTURE TECHNIQUE")
h5_run.bold = True
h5_run.font.size = Pt(18)
h5_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

h5_1 = doc.add_paragraph()
h5_1_run = h5_1.add_run("5.1 Stack technologique recommandée")
h5_1_run.bold = True
h5_1_run.font.size = Pt(14)
h5_1_run.font.color.rgb = RGBColor(255, 102, 0)

table7 = doc.add_table(rows=8, cols=3)
table7.style = 'Table Grid'
headers7 = ["Couche", "Technologie", "Justification"]
for i, header in enumerate(headers7):
    table7.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

stack = [
    ("Frontend", "React.js / Next.js", "Performance, SEO, maintenabilité"),
    ("Backend", "Node.js / Django / Laravel", "Rapidité de développement, robustesse"),
    ("Base de données", "PostgreSQL", "Fiabilité, performances, open source"),
    ("Hébergement", "AWS / Azure / OVH", "Scalabilité, disponibilité 99.9%"),
    ("CDN", "CloudFlare", "Performance, sécurité DDoS"),
    ("SMS Gateway", "Orange SMS API / Local provider", "Couverture nationale"),
    ("Authentification", "JWT + OAuth2", "Sécurité, standard industriel"),
]

for i, (couche, tech, just) in enumerate(stack, 1):
    table7.rows[i].cells[0].paragraphs[0].add_run(couche).font.size = Pt(10)
    table7.rows[i].cells[1].paragraphs[0].add_run(tech).font.size = Pt(10)
    table7.rows[i].cells[2].paragraphs[0].add_run(just).font.size = Pt(10)

set_table_style(table7)

doc.add_paragraph()

h5_2 = doc.add_paragraph()
h5_2_run = h5_2.add_run("5.2 Exigences non-fonctionnelles")
h5_2_run.bold = True
h5_2_run.font.size = Pt(14)
h5_2_run.font.color.rgb = RGBColor(255, 102, 0)

p = doc.add_paragraph()
run_perf = p.add_run("Performance :")
run_perf.bold = True
run_perf.font.size = Pt(11)

perf_items = [
    "Temps de chargement < 3 secondes",
    "Support de 1000 utilisateurs simultanés",
    "Disponibilité 99.9%"
]

for item in perf_items:
    p = doc.add_paragraph()
    p.add_run(f"• {item}").font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(1)

p = doc.add_paragraph()
run_sec = p.add_run("Sécurité :")
run_sec.bold = True
run_sec.font.size = Pt(11)

sec_items = [
    "Certificat SSL/TLS obligatoire",
    "Chiffrement des données sensibles",
    "Protection contre injections SQL et XSS",
    "Conformité RGPD (protection des données personnelles)"
]

for item in sec_items:
    p = doc.add_paragraph()
    p.add_run(f"• {item}").font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(1)

doc.add_paragraph()

h5_3 = doc.add_paragraph()
h5_3_run = h5_3.add_run("5.3 Intégrations envisagées")
h5_3_run.bold = True
h5_3_run.font.size = Pt(14)
h5_3_run.font.color.rgb = RGBColor(255, 102, 0)

integrations = [
    "Passerelle SMS pour notifications",
    "Service email (SMTP)",
    "API géolocalisation pour adresses",
    "Système de paiement (phase ultérieure)"
]

for integ in integrations:
    p = doc.add_paragraph()
    p.add_run(f"• {integ}").font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(1)

doc.add_page_break()

# ==================== 6. DESIGN ET UX ====================

h6 = doc.add_paragraph()
h6_run = h6.add_run("6. DESIGN ET EXPÉRIENCE UTILISATEUR")
h6_run.bold = True
h6_run.font.size = Pt(18)
h6_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

h6_1 = doc.add_paragraph()
h6_1_run = h6_1.add_run("6.1 Identité visuelle")
h6_1_run.bold = True
h6_1_run.font.size = Pt(14)
h6_1_run.font.color.rgb = RGBColor(255, 102, 0)

p = doc.add_paragraph()
p.add_run("La plateforme respectera la charte graphique de la CIE :").font.size = Pt(11)

colors = [
    "Orange CIE (couleur principale) : #FF6600",
    "Bleu foncé (couleur secondaire) : #003366",
    "Blanc (fond) : #FFFFFF",
    "Gris (textes) : #333333"
]

for color in colors:
    p = doc.add_paragraph()
    p.add_run(f"• {color}").font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(1)

doc.add_paragraph()

h6_2 = doc.add_paragraph()
h6_2_run = h6_2.add_run("6.2 Principes UX")
h6_2_run.bold = True
h6_2_run.font.size = Pt(14)
h6_2_run.font.color.rgb = RGBColor(255, 102, 0)

ux_principles = [
    "Simplicité : navigation intuitive en 3 clics maximum",
    "Accessibilité : conformité WCAG 2.1 niveau AA",
    "Responsive : adaptation mobile, tablette, desktop",
    "Feedback : messages de confirmation et d'erreur clairs",
    "Performance : optimisation pour connexions 3G"
]

for principle in ux_principles:
    p = doc.add_paragraph()
    p.add_run(f"• {principle}").font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(1)

doc.add_paragraph()

h6_3 = doc.add_paragraph()
h6_3_run = h6_3.add_run("6.3 Livrables design")
h6_3_run.bold = True
h6_3_run.font.size = Pt(14)
h6_3_run.font.color.rgb = RGBColor(255, 102, 0)

design_liv = [
    "Maquettes haute fidélité (Figma/Adobe XD)",
    "Prototypes interactifs",
    "Guide de style complet",
    "Kit UI components"
]

for liv in design_liv:
    p = doc.add_paragraph()
    p.add_run(f"• {liv}").font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(1)

doc.add_page_break()

# ==================== 7. PLANIFICATION ====================

h7 = doc.add_paragraph()
h7_run = h7.add_run("7. PLANIFICATION")
h7_run.bold = True
h7_run.font.size = Pt(18)
h7_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

h7_1 = doc.add_paragraph()
h7_1_run = h7_1.add_run("7.1 Phases du projet")
h7_1_run.bold = True
h7_1_run.font.size = Pt(14)
h7_1_run.font.color.rgb = RGBColor(255, 102, 0)

table8 = doc.add_table(rows=7, cols=4)
table8.style = 'Table Grid'
headers8 = ["Phase", "Activités", "Durée", "Livrables"]
for i, header in enumerate(headers8):
    table8.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

phases = [
    ("1. Conception", "Analyse besoins, specs fonctionnelles", "3 semaines", "Cahier des charges validé"),
    ("2. Design", "Maquettes, prototypes", "4 semaines", "Maquettes validées"),
    ("3. Développement", "Frontend + Backend + Base données", "8 semaines", "Application fonctionnelle"),
    ("4. Intégration", "Connexion API, SMS, Email", "2 semaines", "Intégrations testées"),
    ("5. Tests", "Tests unitaires, intégration, UAT", "3 semaines", "Rapport de tests"),
    ("6. Déploiement", "Mise en production, formation", "2 semaines", "Application en ligne"),
]

for i, (phase, activites, duree, livrables) in enumerate(phases, 1):
    table8.rows[i].cells[0].paragraphs[0].add_run(phase).font.size = Pt(10)
    table8.rows[i].cells[1].paragraphs[0].add_run(activites).font.size = Pt(10)
    table8.rows[i].cells[2].paragraphs[0].add_run(duree).font.size = Pt(10)
    table8.rows[i].cells[3].paragraphs[0].add_run(livrables).font.size = Pt(10)

set_table_style(table8)

doc.add_paragraph()

h7_2 = doc.add_paragraph()
h7_2_run = h7_2.add_run("7.2 Durée totale estimée")
h7_2_run.bold = True
h7_2_run.font.size = Pt(14)
h7_2_run.font.color.rgb = RGBColor(255, 102, 0)

p = doc.add_paragraph()
run_duree = p.add_run("Durée totale : 22 semaines (environ 5-6 mois)")
run_duree.bold = True
run_duree.font.size = Pt(11)

doc.add_paragraph()

h7_3 = doc.add_paragraph()
h7_3_run = h7_3.add_run("7.3 Jalons clés")
h7_3_run.bold = True
h7_3_run.font.size = Pt(14)
h7_3_run.font.color.rgb = RGBColor(255, 102, 0)

jalons = [
    "J+21 : Validation cahier des charges",
    "J+49 : Validation maquettes",
    "J+105 : Fin développement",
    "J+133 : Recette interne",
    "J+154 : Mise en production"
]

for jalon in jalons:
    p = doc.add_paragraph()
    p.add_run(f"• {jalon}").font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(1)

doc.add_page_break()

# ==================== 8. BUDGET PRÉVISIONNEL ====================

h8 = doc.add_paragraph()
h8_run = h8.add_run("8. BUDGET PRÉVISIONNEL")
h8_run.bold = True
h8_run.font.size = Pt(18)
h8_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

h8_1 = doc.add_paragraph()
h8_1_run = h8_1.add_run("8.1 Estimation des coûts")
h8_1_run.bold = True
h8_1_run.font.size = Pt(14)
h8_1_run.font.color.rgb = RGBColor(255, 102, 0)

table9 = doc.add_table(rows=9, cols=3)
table9.style = 'Table Grid'
headers9 = ["Poste", "Description", "Budget estimé (FCFA)"]
for i, header in enumerate(headers9):
    table9.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

budget = [
    ("Conception", "Analyse, spécifications", "2 000 000"),
    ("Design UI/UX", "Maquettes, prototypes", "3 000 000"),
    ("Développement Frontend", "Interface utilisateur", "6 000 000"),
    ("Développement Backend", "API, logique métier", "7 000 000"),
    ("Base de données", "Conception, optimisation", "1 500 000"),
    ("Tests & Qualité", "Tests, recette", "2 000 000"),
    ("Déploiement", "Mise en production", "1 500 000"),
    ("TOTAL", "", "23 000 000"),
]

for i, (poste, desc, montant) in enumerate(budget, 1):
    if i < 8:
        table9.rows[i].cells[0].paragraphs[0].add_run(poste).font.size = Pt(10)
        table9.rows[i].cells[1].paragraphs[0].add_run(desc).font.size = Pt(10)
        table9.rows[i].cells[2].paragraphs[0].add_run(montant).font.size = Pt(10)
    else:
        for j, val in enumerate([poste, desc, montant]):
            run = table9.rows[i].cells[j].paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            run.bold = True

set_table_style(table9)

doc.add_paragraph()

h8_2 = doc.add_paragraph()
h8_2_run = h8_2.add_run("8.2 Coûts récurrents (annuels)")
h8_2_run.bold = True
h8_2_run.font.size = Pt(14)
h8_2_run.font.color.rgb = RGBColor(255, 102, 0)

table10 = doc.add_table(rows=5, cols=3)
table10.style = 'Table Grid'
headers10 = ["Poste", "Description", "Coût annuel (FCFA)"]
for i, header in enumerate(headers10):
    table10.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

recurring = [
    ("Hébergement", "Serveur cloud, CDN", "1 200 000"),
    ("Noms de domaine", "cie-services.ci", "50 000"),
    ("SMS", "Passerelle notification", "500 000"),
    ("Maintenance", "Correctifs, mises à jour", "3 000 000"),
]

for i, (poste, desc, cout) in enumerate(recurring, 1):
    table10.rows[i].cells[0].paragraphs[0].add_run(poste).font.size = Pt(10)
    table10.rows[i].cells[1].paragraphs[0].add_run(desc).font.size = Pt(10)
    table10.rows[i].cells[2].paragraphs[0].add_run(cout).font.size = Pt(10)

set_table_style(table10)

doc.add_page_break()

# ==================== 9. ÉQUIPE PROJET ====================

h9 = doc.add_paragraph()
h9_run = h9.add_run("9. ÉQUIPE PROJET")
h9_run.bold = True
h9_run.font.size = Pt(18)
h9_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

h9_1 = doc.add_paragraph()
h9_1_run = h9_1.add_run("9.1 Équipe de développement")
h9_1_run.bold = True
h9_1_run.font.size = Pt(14)
h9_1_run.font.color.rgb = RGBColor(255, 102, 0)

table11 = doc.add_table(rows=6, cols=3)
table11.style = 'Table Grid'
headers11 = ["Rôle", "Responsabilités", "Charge"]
for i, header in enumerate(headers11):
    table11.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

team = [
    ("Chef de projet", "Coordination, suivi, reporting", "Temps plein"),
    ("Designer UI/UX", "Maquettes, expérience utilisateur", "Temps partiel"),
    ("Développeur Frontend", "Interface utilisateur", "Temps plein"),
    ("Développeur Backend", "API, logique métier", "Temps plein"),
    ("Testeur QA", "Qualité, recette", "Temps partiel"),
]

for i, (role, resp, charge) in enumerate(team, 1):
    table11.rows[i].cells[0].paragraphs[0].add_run(role).font.size = Pt(10)
    table11.rows[i].cells[1].paragraphs[0].add_run(resp).font.size = Pt(10)
    table11.rows[i].cells[2].paragraphs[0].add_run(charge).font.size = Pt(10)

set_table_style(table11)

doc.add_paragraph()

h9_2 = doc.add_paragraph()
h9_2_run = h9_2.add_run("9.2 Partie CIE (validation et pilotage)")
h9_2_run.bold = True
h9_2_run.font.size = Pt(14)
h9_2_run.font.color.rgb = RGBColor(255, 102, 0)

cie_team = [
    "Directeur de projet (validation stratégique)",
    "Responsable MOA (validation fonctionnelle)",
    "Représentant technique (coordinateur SI)",
    "Utilisateurs clés (tests et recette)"
]

for member in cie_team:
    p = doc.add_paragraph()
    p.add_run(f"• {member}").font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(1)

doc.add_page_break()

# ==================== 10. RISQUES ET MITIGATION ====================

h10 = doc.add_paragraph()
h10_run = h10.add_run("10. RISQUES ET MITIGATION")
h10_run.bold = True
h10_run.font.size = Pt(18)
h10_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

table12 = doc.add_table(rows=6, cols=4)
table12.style = 'Table Grid'
headers12 = ["Risque", "Probabilité", "Impact", "Mitigation"]
for i, header in enumerate(headers12):
    table12.rows[0].cells[i].paragraphs[0].add_run(header).bold = True

risks = [
    ("Retard de validation", "Moyenne", "Élevé", "Planifier les revues à l'avance"),
    ("Évolutions des besoins", "Élevée", "Moyen", "Méthodologie agile, sprints courts"),
    ("Problèmes techniques", "Faible", "Élevé", "Architecture robuste, tests continus"),
    ("Résistance au changement", "Moyenne", "Moyen", "Formation, accompagnement"),
    ("Faille de sécurité", "Faible", "Critique", "Audits sécurité, normes OWASP"),
]

for i, (risque, prob, impact, mitigation) in enumerate(risks, 1):
    table12.rows[i].cells[0].paragraphs[0].add_run(risque).font.size = Pt(10)
    table12.rows[i].cells[1].paragraphs[0].add_run(prob).font.size = Pt(10)
    table12.rows[i].cells[2].paragraphs[0].add_run(impact).font.size = Pt(10)
    table12.rows[i].cells[3].paragraphs[0].add_run(mitigation).font.size = Pt(10)

set_table_style(table12)

doc.add_page_break()

# ==================== 11. ANNEXES ====================

h11 = doc.add_paragraph()
h11_run = h11.add_run("11. ANNEXES")
h11_run.bold = True
h11_run.font.size = Pt(18)
h11_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

h11_1 = doc.add_paragraph()
h11_1_run = h11_1.add_run("Annexe A : Glossaire")
h11_1_run.bold = True
h11_1_run.font.size = Pt(14)
h11_1_run.font.color.rgb = RGBColor(255, 102, 0)

glossary = [
    ("CIE", "Compagnie Ivoirienne d'Electricité"),
    ("PEPT", "Programme Électricité Pour Tous"),
    ("UI/UX", "User Interface / User Experience"),
    ("API", "Application Programming Interface"),
    ("RGPD", "Règlement Général sur la Protection des Données"),
    ("WCAG", "Web Content Accessibility Guidelines"),
    ("JWT", "JSON Web Token"),
    ("CDN", "Content Delivery Network"),
]

for abbr, definition in glossary:
    p = doc.add_paragraph()
    run_abbr = p.add_run(f"{abbr} : ")
    run_abbr.bold = True
    run_abbr.font.size = Pt(11)
    p.add_run(definition).font.size = Pt(11)

doc.add_paragraph()

h11_2 = doc.add_paragraph()
h11_2_run = h11_2.add_run("Annexe B : Références")
h11_2_run.bold = True
h11_2_run.font.size = Pt(14)
h11_2_run.font.color.rgb = RGBColor(255, 102, 0)

refs = [
    "Site officiel CIE : https://www.cie.ci/",
    "Programme Électricité Pour Tous - Brochure officielle",
    "Charte graphique CIE",
    "Normes WCAG 2.1 - Accessibilité web"
]

for ref in refs:
    p = doc.add_paragraph()
    p.add_run(f"• {ref}").font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(1)

doc.add_paragraph()

h11_3 = doc.add_paragraph()
h11_3_run = h11_3.add_run("Annexe C : Contacts")
h11_3_run.bold = True
h11_3_run.font.size = Pt(14)
h11_3_run.font.color.rgb = RGBColor(255, 102, 0)

p = doc.add_paragraph()
p.add_run("Pour toute question relative à ce cahier des charges, merci de contacter :").font.size = Pt(11)

doc.add_paragraph()

contacts = [
    "Direction des Systèmes d'Information - CIE",
    "Email : dsi@cie.ci",
    "Téléphone : +225 XX XX XX XX",
    "Adresse : Abidjan, Côte d'Ivoire"
]

for contact in contacts:
    p = doc.add_paragraph()
    p.add_run(f"• {contact}").font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(1)

# Pied de page final
doc.add_paragraph()
doc.add_paragraph()

footer_final = doc.add_paragraph()
footer_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_footer = footer_final.add_run("— Fin du document —")
run_footer.italic = True
run_footer.font.size = Pt(10)

# Sauvegarder le document
doc.save('Cahier_des_Charges_CIE_Plateforme_Digitale.docx')
print("Document Word créé avec succès!")